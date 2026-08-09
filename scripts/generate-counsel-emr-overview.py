"""Generate Patient Vault EMR overview Word doc for counsel."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

OUT_DIR = Path(r"C:\Users\Firas\Desktop")
OUT_DOCX = OUT_DIR / "Patient_Vault_EMR_Overview_for_Counsel.docx"
REPO_COPY = Path(r"C:\Users\Firas\patient-vault\Patient_Vault_EMR_Overview_for_Counsel.docx")


def set_run_font(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, size=11, bold=False, italic=False, space_after=8, space_before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 13, 2: 12}
    space_before = {1: 16, 2: 12}
    return add_para(
        doc,
        text,
        size=sizes.get(level, 11),
        bold=True,
        space_before=space_before.get(level, 10),
        space_after=6,
    )


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, size=10)
    doc.add_paragraph()


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_para(
        doc,
        "CONFIDENTIAL — FOR LEGAL COUNSEL REVIEW",
        size=10,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        "PATIENT VAULT (MODERN MEDICINE)",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        "EMR Overview for Healthcare Counsel",
        size=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18,
    )

    add_para(doc, "To: Ayesha Mehdi, Esq.", size=11, space_after=2)
    add_para(doc, "From: Firas Khamis, MD / Modern Medicine", size=11, space_after=2)
    add_para(doc, "Re: Product overview for HIPAA / healthcare counsel review", size=11, space_after=2)
    add_para(doc, "Production URL: https://app.patientvault.care", size=11, space_after=2)
    add_para(doc, "Date: August 3, 2026", size=11, space_after=12)

    add_para(
        doc,
        "This memorandum describes the electronic medical record (EMR) system for counsel review. "
        "It is not a legal opinion, not a final HIPAA determination, and not adopted policy. "
        "Please treat access credentials and any demo PHI as confidential.",
        italic=True,
        space_after=12,
    )

    add_heading(doc, "1. What the product is")
    add_para(
        doc,
        "Patient Vault is a cloud-hosted clinical EMR used by Modern Medicine for ambulatory care. "
        "Clinicians and authorized staff use it in a web browser (desktop or phone) to:",
    )
    add_bullets(
        doc,
        [
            "Maintain patient charts (demographics, history, diagnoses, medications, allergies, social history, and related clinical sections)",
            "Create, edit, sign, and revise encounter notes",
            "Schedule clinic visits and manage encounters",
            "Upload and store clinical documents",
            "Place and track clinical orders",
            "Send internal clinic messages and set patient-related reminders",
            "Use optional AI-assisted drafting tools (speech-to-text and HPI drafting) that run on AWS services intended for use under an AWS Business Associate Addendum (BAA)",
        ],
    )
    add_para(doc, "The live application address is https://app.patientvault.care.", space_after=8)

    add_heading(doc, "2. Who uses it and how access works")
    add_table(
        doc,
        ["Item", "Current design"],
        [
            ["Access", "Unique user accounts (email + password) over HTTPS"],
            ["Roles", "ADMIN, CLINICIAN, STAFF, READONLY"],
            ["Session", "Server-side sessions with idle timeout (short in production)"],
            ["MFA", "Time-based one-time passwords (TOTP) available / in use for workforce accounts"],
            [
                "Audit",
                "Application audit log for login, PHI access/create/update/delete, AI use, and config changes",
            ],
            [
                "Counsel account",
                "Read-only role recommended for review (view, not edit clinical data)",
            ],
        ],
    )
    add_para(
        doc,
        "Your review account is configured as READONLY so you can inspect the product and workflows without changing charts.",
    )

    add_heading(doc, "3. What ePHI the system handles")
    add_para(doc, "Reasonably anticipated ePHI categories include:")
    add_bullets(
        doc,
        [
            "Patient identifiers and demographics (name, DOB, MRN, contact, insurance fields as entered)",
            "Clinical chart content and notes",
            "Diagnoses, medications, allergies, history",
            "Uploaded documents (e.g., reports, PDFs, visit recordings when used)",
            "Scheduling / encounter metadata",
            "Internal care messaging and reminders as used by the clinic",
        ],
    )
    add_para(
        doc,
        'Personal "scratch" notes and certain AI Listen text saves are stored as private per-user content '
        "(not shared across users as part of the shared chart), but they may still contain PHI and should be treated as such.",
    )

    add_heading(doc, "4. Technical environment (high level)")
    add_table(
        doc,
        ["Component", "Provider / notes"],
        [
            ["Application", "Next.js web app on AWS Lightsail"],
            ["Database", "AWS RDS PostgreSQL (encrypted storage path; TLS to DB)"],
            [
                "Documents",
                "Local volume and/or AWS S3 (SSE), depending on production configuration",
            ],
            ["TLS", "HTTPS via Caddy on the application host"],
            ["AI transcription", "Amazon Transcribe Medical"],
            [
                "AI text drafting",
                "Amazon Bedrock (Claude models), under clinic AWS account / BAA posture",
            ],
            [
                "Fax",
                "Present in product; production fax vendor must be BAA-covered before live PHI faxing",
            ],
        ],
    )
    add_para(
        doc,
        "Application-layer field encryption is used for selected sensitive patient fields; note content encryption "
        "helpers exist in the codebase. Exact field coverage and key management should be confirmed against the "
        "Security Risk Analysis and current production configuration during your review.",
    )

    add_heading(doc, "5. Clinical / product features counsel may want to walk through")
    add_numbered(
        doc,
        [
            "Patient list & chart — open a patient; left-rail clinical sections; notes panel",
            "Notes — draft, sign, revise; PDF export where available",
            "Schedule — clinic day schedule and visit association",
            "Documents — upload/view clinical files in the chart",
            "Orders — create/review orders tied to care",
            "Messages / Reminders / Contacts — clinic operations tools",
            "Users & security (admin) — roles, password reset, unlock, MFA settings",
            "AI Listen / Visit recorder — record visit audio to transcript + HPI draft (AWS Transcribe Medical + Bedrock); audio may be stored to the chart when using visit recorder",
            "Ask AI / guidelines assists — clinician-facing drafting aids (not a substitute for clinical judgment)",
        ],
    )

    add_heading(doc, "6. Documents already prepared for your review packet")
    add_para(doc, "Please also receive (or request) the following drafts already prepared for counsel:")
    add_bullets(
        doc,
        [
            "Security Risk Analysis draft",
            "Legal / policy counsel review draft Word document",
            "Production / deploy operational notes as needed for BAA and hosting questions",
        ],
    )
    add_para(
        doc,
        "Counsel should confirm whether these drafts meet HIPAA Privacy/Security Rule expectations for this practice, "
        "Nevada/state requirements, notice of privacy practices, BAAs with AWS and any other vendors, and workforce/device policies.",
    )

    add_heading(doc, "7. How to sign in (your account)")
    add_para(doc, "URL: https://app.patientvault.care/login")
    add_para(
        doc,
        "Credentials will be provided separately in the cover email (username/email + temporary password). "
        "On first login you may be required to set a new password. Please:",
    )
    add_bullets(
        doc,
        [
            "Do not share the password",
            "Enable MFA if prompted / available for your account",
            "Use the account only for professional review",
            "Avoid entering real patient data unless this environment is confirmed appropriate for that purpose",
        ],
    )
    add_para(
        doc,
        "If login fails, contact Dr. Khamis; production outages are usually infrastructure (database connectivity), not account policy.",
    )

    add_heading(doc, "8. Open items we specifically want counsel guidance on")
    add_numbered(
        doc,
        [
            "Adequacy of current policies, NPP, and BAAs (especially AWS; fax/AI vendors if enabled for PHI)",
            "Whether READONLY counsel access and audit logging are appropriate for outside counsel review",
            "Retention, amendment, and patient-rights workflows as implemented vs. required",
            "Use of AI transcription/drafting under BAA and notice/consent expectations",
            'Password-free "test mode" visit recorder (if ever enabled) — must remain off for real clinic use',
            "Any Nevada-specific telehealth, privacy, or medical-board considerations for this EMR",
        ],
    )

    add_heading(doc, "9. Contact")
    add_para(doc, "Practice / Security Officer: Firas Khamis, MD", space_after=2)
    add_para(doc, "System: Patient Vault / Modern Medicine", space_after=2)
    add_para(doc, "Production: https://app.patientvault.care", space_after=12)
    add_para(doc, "Thank you for reviewing this system.")

    doc.save(OUT_DOCX)
    doc.save(REPO_COPY)
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {REPO_COPY}")


if __name__ == "__main__":
    main()
