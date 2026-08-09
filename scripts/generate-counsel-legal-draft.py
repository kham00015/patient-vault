"""Generate attorney-facing Patient Vault legal policy review packet."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_DIR = Path(r"C:\Users\Firas\Desktop")
OUT_DOCX = OUT_DIR / "Patient_Vault_Legal_Policy_Counsel_Review_Draft.docx"
OUT_HTML = OUT_DIR / "Patient_Vault_Legal_Policy_Counsel_Review_Draft.html"
REPO_MD = Path(r"C:\Users\Firas\patient-vault\LEGAL_POLICY_COUNSEL_REVIEW_DRAFT.md")


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, *, size=11, bold=False, italic=False, space_after=8, space_before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    # Manual headings for consistent Times New Roman look
    sizes = {0: 16, 1: 13, 2: 12, 3: 11}
    space_before = {0: 6, 1: 16, 2: 12, 3: 10}
    p = add_para(
        doc,
        text,
        size=sizes.get(level, 11),
        bold=True,
        space_before=space_before.get(level, 10),
        space_after=6,
    )
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=11)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Cover
    add_para(doc, "CONFIDENTIAL — ATTORNEY WORK PRODUCT CANDIDATE", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "DRAFT FOR LEGAL COUNSEL REVIEW ONLY", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "PATIENT VAULT", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(
        doc,
        "HIPAA Privacy, Security, and Operational Policy Packet",
        size=14,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        "Draft Policies, Procedures, and Counsel Review Checklist",
        size=12,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18,
    )

    add_para(doc, f"Document status: Draft — not adopted", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, f"Version date: {date.today().isoformat()}", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Intended recipient: Outside healthcare counsel", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Application: Patient Vault (https://app.patientvault.care)", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    add_para(
        doc,
        "IMPORTANT NOTICE: This packet is a working draft prepared to facilitate review by qualified healthcare counsel. "
        "It is not legal advice, is not a final compliance program, and must not be published to patients, workforce members, "
        "vendors, payers, or the public until reviewed, revised as needed, and formally adopted by the Organization. "
        "Bracketed placeholders (example: [Legal Entity Name]) require completion by counsel or compliance leadership.",
        size=10,
        italic=True,
        space_after=18,
    )

    # Transmittal
    add_heading(doc, "I. Transmittal to Counsel", 1)
    add_para(
        doc,
        "This packet is submitted for independent legal review in connection with the Organization’s operation and use of "
        "Patient Vault, a cloud-hosted electronic medical record and clinic operations application that may create, receive, "
        "maintain, or transmit electronic protected health information (“ePHI”).",
    )
    add_para(doc, "Counsel is requested to:")
    add_numbered(
        doc,
        [
            "Confirm the correct legal entity or entities, ownership structure, and HIPAA role(s) (covered entity, business associate, or hybrid).",
            "Revise these draft policies into final internal policies, workforce procedures, patient-facing notices, and vendor/customer contract terms, as appropriate.",
            "Insert applicable state retention, breach-notification, telehealth, medical-board, prescribing, and consumer-health privacy requirements.",
            "Identify gaps between the described technical controls and the administrative, physical, and technical safeguards required by the HIPAA Security Rule.",
            "Advise whether additional consent, notice, or documentation is required for AI-assisted documentation, ambient listening/transcription, faxing, messaging, and any future e-prescribing features.",
            "Return a marked-up draft and a short list of required pre-go-live legal actions.",
        ],
    )

    add_heading(doc, "II. Defined Terms", 1)
    add_para(
        doc,
        "For purposes of this draft, the following terms have the meanings set forth below. Counsel may revise definitions to "
        "align with final policy language and applicable law.",
    )
    add_bullets(
        doc,
        [
            "“Organization” means [Legal Entity Name], including any affiliated practice entities designated by counsel.",
            "“Patient Vault” means the web application hosted at https://app.patientvault.care and related production infrastructure.",
            "“PHI” and “ePHI” have the meanings assigned under HIPAA and its implementing regulations.",
            "“Workforce” means employees, contractors, clinicians, and other persons under the Organization’s direct control who use Patient Vault.",
            "“Privacy Officer” means the individual designated under Section IV.",
            "“Security Officer” means the individual designated under Section IV.",
            "“Approved Vendor” means a vendor that has executed an appropriate business associate agreement (“BAA”) where required and has been approved for PHI processing.",
        ],
    )

    add_heading(doc, "III. Issues for Counsel Determination", 1)
    add_para(doc, "Before adoption or production use with real patient data, counsel should resolve at least the following:")
    add_numbered(
        doc,
        [
            "Is Patient Vault operated solely by a covered entity medical practice, by a separate software entity, or both?",
            "Which legal entity owns Patient Vault software and which legal entity owns the clinical practice using it?",
            "If separate entities are involved, is a BAA required between them, and in which direction?",
            "Has the AWS Business Associate Addendum been accepted in AWS Artifact for the production AWS account?",
            "Are all AWS services used for PHI HIPAA-eligible under the AWS BAA as configured?",
            "What state medical-record retention periods apply to the Organization’s specialty and locations?",
            "What state privacy, breach-notification, and consumer health data laws apply?",
            "What consent or notice is required for AI-assisted documentation, ambient capture, and Amazon Transcribe Medical workflows?",
            "What consent is required for SMS, email, fax, or patient messaging features?",
            "What prescribing, EPCS, DEA, NPI, and pharmacy-network requirements apply before any e-prescribing launch?",
            "What Notice of Privacy Practices, website terms, and patient consent forms must be displayed or collected?",
            "What audit-log and policy-document retention periods are required or recommended?",
            "What cyber insurance, incident-response counsel, and breach-notification vendors should be designated?",
        ],
    )

    add_heading(doc, "IV. Governance", 1)
    add_heading(doc, "A. Privacy Officer", 2)
    add_para(
        doc,
        "The Organization shall designate a Privacy Officer responsible for HIPAA Privacy Rule compliance, patient rights "
        "workflows, use/disclosure review, privacy training, privacy incident evaluation, and maintenance of privacy policies.",
    )
    add_para(doc, "Designated Privacy Officer: [Name / Title]")
    add_heading(doc, "B. Security Officer", 2)
    add_para(
        doc,
        "The Organization shall designate a Security Officer responsible for HIPAA Security Rule compliance, risk analysis "
        "and risk management, technical safeguards, access control oversight, security incident response, vendor security "
        "review, and backup/disaster-recovery oversight.",
    )
    add_para(doc, "Designated Security Officer: [Name / Title]")
    add_heading(doc, "C. Policy Review Cadence", 2)
    add_para(doc, "These policies shall be reviewed at least annually and upon any of the following:")
    add_bullets(
        doc,
        [
            "A material security or privacy incident",
            "A material system architecture change",
            "Onboarding of a new vendor that creates, receives, maintains, or transmits PHI",
            "Launch of AI, transcription, fax, messaging, or e-prescribing features",
        ],
    )

    add_heading(doc, "V. System Description for Counsel", 1)
    add_para(
        doc,
        "Patient Vault is a web-based clinical application supporting patient registration, scheduling, charting, clinical "
        "notes (including signed-note revision with audit stamps), document storage, encounters, orders, internal messaging, "
        "forms, reminders, contacts, private clinician notes, and administrative user management.",
    )
    add_para(
        doc,
        "Production hosting currently includes AWS Lightsail for the application, AWS RDS PostgreSQL for the database, and "
        "AWS S3 for document storage. Documents in S3 are not publicly accessible; public access is blocked, versioning is "
        "enabled, and server-side encryption is enabled. Database backups are automated and encrypted.",
    )
    add_para(
        doc,
        "Current technical controls include unique user accounts, role-based access control, password policy enforcement, "
        "multi-factor authentication support, account lockout after repeated failed logins, audit logging, secure production "
        "cookies, a five-minute idle session timeout, HTTPS, security headers, and login rate limiting.",
    )
    add_para(
        doc,
        "Current AI and transcription capabilities, where enabled, use Amazon Bedrock and Amazon Transcribe Medical under "
        "the Organization’s AWS account and BAA posture. No vendor that receives PHI shall be used unless an appropriate BAA "
        "is in place and the service is configured for HIPAA-eligible use. Consumer AI tools are prohibited for PHI.",
    )

    add_heading(doc, "VI. Data Classification and PHI Handling", 1)
    add_para(
        doc,
        "Patient Vault may contain PHI including names, demographics, contact information, medical record numbers, "
        "appointments, clinical notes, diagnoses, medications, uploaded records, referral documents, forms, care-related "
        "messages, and audit records that identify who accessed patient information.",
    )
    add_para(
        doc,
        "Workforce members shall treat chart content as confidential. PHI shall not be copied to personal email, personal "
        "cloud storage, consumer AI tools, personal devices, unapproved messaging applications, or unapproved external "
        "services. Access is limited to treatment, payment, healthcare operations, or other legally permitted purposes.",
    )
    add_para(doc, "Data shall be classified at least as follows:")
    add_bullets(
        doc,
        [
            "PHI / ePHI — highest protection",
            "Confidential business information — contracts, security procedures, pricing, internal policies",
            "Public / low-risk information — non-identifying public website content",
        ],
    )
    add_para(
        doc,
        "Developers and administrators shall not place PHI in source code, public repositories, issue trackers, screenshots, "
        "unsecured support channels, or unnecessary debugging output. Troubleshooting should prefer resource identifiers and "
        "non-PHI technical details.",
    )

    add_heading(doc, "VII. HIPAA Privacy Rule Draft Policy", 1)
    add_para(
        doc,
        "The Organization may use and disclose PHI only as permitted by HIPAA, applicable state law, and Organizational "
        "policy. Routine uses may include treatment, payment, healthcare operations, patient access requests, care "
        "coordination, quality review, scheduling, and administrative operations. Uses outside those categories require "
        "Privacy Officer or counsel review before disclosure.",
    )
    add_para(
        doc,
        "The minimum necessary standard applies to most uses and disclosures. Workforce members should access only the "
        "information needed for assigned duties. Clinicians may require broader access for treatment; other roles receive "
        "least-privilege access. Role assignments remain a management responsibility.",
    )
    add_para(
        doc,
        "The Organization shall maintain a Notice of Privacy Practices approved by counsel. Patients may have rights of "
        "access, amendment, restriction, confidential communications, accounting of disclosures, and complaint without "
        "retaliation. Patient Vault can support retrieval of charts, documents, and audit logs, but legal procedures and "
        "deadlines must be defined by counsel.",
    )

    add_heading(doc, "VIII. HIPAA Security Rule Draft Policies", 1)
    add_heading(doc, "A. Administrative Safeguards", 2)
    add_para(
        doc,
        "The Security Officer shall maintain a written risk analysis covering hosting, database, document storage, user "
        "access, remote access, endpoints, backups, vendors, audit logs, incident response, AI, fax, transcription, and "
        "e-prescribing. A risk-management plan shall document how identified risks are addressed through technical and "
        "administrative controls. Access grant/change/removal procedures, sanctions, contingency planning, and workforce "
        "training are required elements.",
    )
    add_heading(doc, "B. Physical Safeguards", 2)
    add_para(
        doc,
        "Although Patient Vault is cloud-hosted, physical safeguards remain required for workstations and facilities. "
        "Screens shall be protected from unauthorized viewing; workstations shall lock when unattended; printed PHI shall "
        "be minimized and secured; and facility access shall be controlled.",
    )
    add_heading(doc, "C. Technical Safeguards", 2)
    add_para(
        doc,
        "Unique user identification, authentication (password plus MFA where required), automatic logoff, encryption in "
        "transit (HTTPS), encryption of stored PHI where implemented, and audit controls are required. Shared accounts for "
        "routine operations are prohibited. Emergency access, if any, must be defined, logged, and reviewed.",
    )

    add_heading(doc, "IX. Access Management Procedure (Draft)", 1)
    add_para(
        doc,
        "Before account creation, management shall verify identity, role, employment/contractor status, and need for access. "
        "Least-privilege roles shall be assigned. Temporary passwords require change at first login. MFA should be enabled "
        "before access to real patient charts. Access reviews shall occur at least quarterly. Departure or role change "
        "requires prompt disablement and session revocation.",
    )

    add_heading(doc, "X. Audit Logging and Review (Draft)", 1)
    add_para(
        doc,
        "Audit logs should record PHI access, note creation/update/sign/revision, document upload/download, login and failed "
        "login events, MFA-related events, user-administration events, AI queries, and exports such as printable note views. "
        "Logs should include user identifiers, patient identifiers, resource identifiers, timestamps, and IP/user-agent where "
        "available, while avoiding unnecessary raw clinical text in metadata.",
    )
    add_para(
        doc,
        "Audit logs shall be reviewed after suspected inappropriate access, after incidents, before certain patient-rights "
        "responses, and during periodic compliance checks. Reviews shall be documented.",
    )

    add_heading(doc, "XI. Medical Record Integrity; Signed Notes and Revisions (Draft)", 1)
    add_para(
        doc,
        "Clinical documentation integrity is material to compliance and malpractice risk. Draft notes may be edited freely "
        "before signature. Upon signature, the note is marked signed with date/time and signing user.",
    )
    add_para(
        doc,
        "Patient Vault permits post-signature revision with compliance stamps. Each revision stores a prior-content snapshot, "
        "records who revised and when, increments a revision counter, and retains initiated/signed/revised stamps on the note "
        "and printable output. Signed notes remain non-deletable through ordinary deletion workflows. Counsel should confirm "
        "whether this model satisfies applicable record-integrity expectations or whether an addendum-only model is preferred.",
    )

    add_heading(doc, "XII. Retention and Destruction (Draft)", 1)
    add_para(
        doc,
        "Retention periods shall be set by counsel based on federal law, state law, payer rules, and specialty standards. "
        "Until counsel confirms otherwise, the Organization shall favor retention and archival over permanent deletion. "
        "Silent deletion of inaccurate records is disfavored; correction, addendum, or controlled revision is preferred. "
        "HIPAA documentation is commonly retained at least six years; medical-record retention may be longer. Destruction, "
        "when permitted, shall be authorized and documented. Backup copies may persist until backup retention expires.",
    )

    add_heading(doc, "XIII. Backup and Disaster Recovery (Draft)", 1)
    add_para(
        doc,
        "Current controls include encrypted automated RDS backups (presently seven-day retention) and encrypted private S3 "
        "document storage with versioning. The Security Officer shall verify backup status at least monthly and after major "
        "changes, and shall conduct a restore test at least annually. Leadership shall define recovery time and recovery "
        "point objectives appropriate to clinic operations.",
    )

    add_heading(doc, "XIV. Incident Response and Breach Notification (Draft)", 1)
    add_para(
        doc,
        "Workforce members must promptly report suspected privacy or security incidents, including lost devices, suspicious "
        "logins, phishing, misdirected faxes, unauthorized chart access, malware/ransomware, accidental disclosure, vendor "
        "incidents, and compromised credentials.",
    )
    add_para(
        doc,
        "Initial response prioritizes containment and evidence preservation. The Privacy Officer and Security Officer shall "
        "assess whether PHI was involved and whether a breach may have occurred. Counsel shall guide risk assessment and "
        "notification analysis. HIPAA and state notification deadlines may apply; state law may be stricter than HIPAA’s "
        "general outer timelines. An incident contact list shall include leadership, counsel, IT, cyber insurance, and key vendors.",
    )

    add_heading(doc, "XV. Vendor and BAA Policy (Draft)", 1)
    add_para(
        doc,
        "No vendor may create, receive, maintain, or transmit PHI on behalf of the Organization until the relationship is "
        "reviewed and, where required, a BAA is executed. AWS shall be addressed through AWS Artifact acceptance of the AWS "
        "Business Associate Addendum for the production account. The BAA does not, by itself, establish compliance; "
        "configuration and use remain Organizational responsibilities.",
    )
    add_para(
        doc,
        "AI, transcription, fax, messaging, and e-prescribing vendors require the same review: covered product tier, "
        "subcontractors, data location, retention, logging, training/use of data, and incident reporting.",
    )

    add_heading(doc, "XVI. Artificial Intelligence Policy (Draft)", 1)
    add_para(
        doc,
        "AI features may assist with drafting, summarization, organization, guidelines review, and ambient HPI drafting. "
        "AI does not replace clinician judgment. AI-generated content that may enter the medical record must be reviewed by "
        "an appropriate human user before reliance, signature, transmission, or clinical decision-making.",
    )
    add_para(
        doc,
        "Workforce members shall not paste PHI into consumer AI tools (including personal ChatGPT accounts, browser "
        "extensions, and unapproved transcription tools). Only Approved Vendor AI integrations under BAA—currently "
        "Amazon Bedrock within the Organization’s AWS environment, where enabled—may process PHI.",
    )
    add_para(
        doc,
        "Approved AI workflows shall follow minimum necessary principles. Outputs are draft until reviewed. Counsel shall "
        "determine whether patient notices, consent language, or special audit labeling are required.",
    )

    add_heading(doc, "XVII. Transcription / Ambient Listening Policy (Draft)", 1)
    add_para(
        doc,
        "Where Amazon Transcribe Medical or similar tools are used, the Organization shall determine whether audio is "
        "temporarily processed only or stored, retention duration, microphone/device controls, and whether state consent "
        "laws require notice or consent for recording clinical conversations. Transcripts are draft until clinician review. "
        "Unnecessary retention of raw audio should be avoided unless counsel approves a retention schedule.",
    )

    add_heading(doc, "XVIII. Fax Policy (Draft)", 1)
    add_para(
        doc,
        "Faxing PHI requires an Approved Vendor under BAA and a written procedure for number verification, cover sheets, "
        "failed-transmission handling, and misdirected-fax escalation. Misdirected faxes are potential privacy incidents.",
    )

    add_heading(doc, "XIX. E-Prescribing Policy (Draft)", 1)
    add_para(
        doc,
        "Patient Vault shall not transmit real prescriptions to pharmacies without an Approved Vendor and counsel-approved "
        "workflow covering identity proofing, DEA/NPI/state license requirements, EPCS for controlled substances, pharmacy "
        "selection, failure handling, and medical-record documentation.",
    )

    add_heading(doc, "XX. Workforce Training and Sanctions (Draft)", 1)
    add_para(
        doc,
        "Training is required before Patient Vault access and at least annually thereafter, covering privacy/security "
        "policies, minimum necessary access, credentials, MFA, incident reporting, workstation security, phishing, and "
        "restrictions on AI/vendors. Training records shall be retained as counsel directs.",
    )
    add_para(
        doc,
        "Violations may result in warning, retraining, access suspension, employment discipline, termination, or reporting "
        "to authorities/licensing boards when required. Sanctions shall be consistent, proportionate, and documented.",
    )

    add_heading(doc, "XXI. Patient-Facing and Commercial Documents Still Needed", 1)
    add_para(doc, "This packet is not a substitute for counsel-approved:")
    add_bullets(
        doc,
        [
            "Notice of Privacy Practices",
            "Website privacy policy and terms of use",
            "Patient consent / communication consent forms",
            "Telehealth consent, if applicable",
            "AI / ambient documentation disclosure language, if required",
            "Medical record request procedures",
        ],
    )
    add_para(
        doc,
        "If Patient Vault is offered to unaffiliated clinics, counsel should also prepare customer terms, subscription "
        "agreement, BAA template, support/SLA terms, data ownership/export/termination assistance, indemnity, and "
        "limitation of liability provisions.",
    )

    add_heading(doc, "XXII. Pre-Adoption Compliance Checklist", 1)
    add_bullets(
        doc,
        [
            "AWS BAA accepted for production account",
            "Unique named accounts for all users; no shared clinical logins",
            "MFA enabled for users with PHI access",
            "Default/demo credentials removed or rotated",
            "Privacy Officer and Security Officer designated in writing",
            "Workforce HIPAA training completed and recorded",
            "Incident-response contacts documented",
            "Backups verified; restore test scheduled",
            "Vendor BAAs completed before AI/fax/transcription/e-prescribing PHI flows",
            "Counsel sign-off on patient notices and go-live legal readiness",
        ],
    )

    add_heading(doc, "Appendix A — Vendor Review Matrix", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["Vendor Category", "Current / Example Services", "BAA Before PHI?", "Counsel Focus"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, size=10, bold=True)

    rows = [
        ("Cloud hosting / DB / object storage", "AWS Lightsail, RDS, S3", "Yes", "Artifact BAA; HIPAA-eligible services; encryption; backups; shared responsibility"),
        ("Clinical AI", "Amazon Bedrock", "Yes (via AWS BAA / eligible services)", "Model access, logging, retention, human review, patient notice"),
        ("Medical transcription", "Amazon Transcribe Medical", "Yes (via AWS BAA / eligible services)", "Recording consent, audio retention, draft review, specialty settings"),
        ("Fax", "[Fax vendor TBD]", "Yes", "Misdirect workflow, logs, transmission security, retention"),
        ("E-prescribing", "[E-Rx vendor TBD]", "Yes", "EPCS, DEA/NPI, identity proofing, pharmacy network, audit trail"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9)

    doc.add_paragraph()
    add_heading(doc, "Appendix B — Approval Block (Do Not Sign Until Counsel Completes Review)", 1)
    add_para(doc, "Reviewed by Legal Counsel: _________________________________ Date: ____________")
    add_para(doc, "Approved by Privacy Officer: ________________________________ Date: ____________")
    add_para(doc, "Approved by Security Officer: _______________________________ Date: ____________")
    add_para(doc, "Approved by Authorized Organizational Officer: ______________ Date: ____________")
    add_para(
        doc,
        "END OF DRAFT — NOT FOR PATIENT, WORKFORCE, OR PUBLIC DISTRIBUTION PENDING COUNSEL APPROVAL.",
        size=10,
        bold=True,
        italic=True,
        space_before=18,
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")

    # Also write a markdown copy into the repo for version control convenience.
    # Keep it brief pointer + note that DOCX is the counsel deliverable.
    REPO_MD.write_text(
        f"""# Patient Vault — Counsel Review Draft

**Deliverable for attorney review:**  
`C:\\Users\\Firas\\Desktop\\Patient_Vault_Legal_Policy_Counsel_Review_Draft.docx`

- Status: Draft for outside healthcare counsel review only  
- Version date: {date.today().isoformat()}  
- Not legal advice; not adopted policy  
- Written in institutional voice for counsel (not addressed to an individual owner)

Regenerate with: `python scripts/generate-counsel-legal-draft.py`
""",
        encoding="utf-8",
    )
    print(f"Wrote {REPO_MD}")


if __name__ == "__main__":
    main()
