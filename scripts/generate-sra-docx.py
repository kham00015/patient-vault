"""Generate Security Risk Analysis Word doc for counsel from SECURITY_RISK_ANALYSIS.md (simplified structured export)."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SRC = Path(r"C:\Users\Firas\patient-vault\SECURITY_RISK_ANALYSIS.md")
OUT = Path(r"C:\Users\Firas\Desktop\Patient_Vault_Security_Risk_Analysis_Draft.docx")
OUT2 = Path(r"C:\Users\Firas\patient-vault\Patient_Vault_Security_Risk_Analysis_Draft.docx")


def set_run_font(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, size=11, bold=False, italic=False, space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    for m in (section.top_margin,):
        pass
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_para(doc, "CONFIDENTIAL — DRAFT FOR LEGAL COUNSEL REVIEW", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "PATIENT VAULT — SECURITY RISK ANALYSIS", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, "Version 0.2 Draft — August 3, 2026 (revises July 26, 2026)", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    for line in md.splitlines():
        s = line.rstrip()
        if not s:
            continue
        if s.startswith("# "):
            continue  # title already set
        if s.startswith("## "):
            add_para(doc, s[3:].strip(), size=13, bold=True, space_before=14, space_after=6)
        elif s.startswith("### "):
            add_para(doc, s[4:].strip(), size=12, bold=True, space_before=10, space_after=4)
        elif s.startswith("> "):
            add_para(doc, s[2:].strip(), size=10, italic=True, space_after=8)
        elif s.startswith("|") and "---" in s:
            continue
        elif s.startswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            add_para(doc, " | ".join(cells), size=9, space_after=2)
        elif s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(s[2:])
            set_run_font(run, size=11)
        elif s.startswith("---"):
            continue
        else:
            # strip simple markdown bold markers for readability
            text = s.replace("**", "")
            add_para(doc, text, size=11)

    doc.save(OUT)
    doc.save(OUT2)
    print(f"Wrote {OUT}")
    print(f"Wrote {OUT2}")


if __name__ == "__main__":
    main()
